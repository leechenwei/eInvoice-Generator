import tkinter as tk
from tkinter import messagebox, simpledialog
from tkinter import ttk
from docx import Document
import os
import json
from tkcalendar import DateEntry

# File paths for company and To Whom data
COMPANY_DATA_FILE = 'companies.json'
TO_WHOM_DATA_FILE = 'to_whom.json'

# Load company data from file
def load_companies():
    if os.path.exists(COMPANY_DATA_FILE):
        with open(COMPANY_DATA_FILE, 'r') as file:
            return json.load(file)
    return {}

# Save company data to file
def save_companies():
    with open(COMPANY_DATA_FILE, 'w') as file:
        json.dump(companies, file)

# Load To Whom data from file
def load_to_whom():
    if os.path.exists(TO_WHOM_DATA_FILE):
        with open(TO_WHOM_DATA_FILE, 'r') as file:
            return json.load(file)
    return []

# Save To Whom data to file
def save_to_whom():
    with open(TO_WHOM_DATA_FILE, 'w') as file:
        json.dump(to_whom_list, file)

def create_invoice(company_name, ssm_number, invoice_number, invoice_date, items, to_whom):
    # Load the template document
    doc = Document('invoice_template.docx')

    # Get the company address and phone number
    company_info = companies[company_name]
    address = company_info.get('address', 'N/A')
    phone_number = company_info.get('phone', 'N/A')

    # Format the phone number
    phone_number = f"{phone_number[:3]}-{phone_number[3:]}"

    # Dictionary for placeholders and their corresponding values
    placeholders = {
        '{CompanyName}': company_name,
        '{SSMNumber}': ssm_number,
        '{PhoneNumber}': phone_number,
        '{InvoiceDate}': invoice_date,
        '{CompanyAddress}': address,
        '{ToWhom}': to_whom,
        '{InvoiceNo}': str(invoice_number)  # Ensure this is a string
    }

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                # Bold the replaced text
                if placeholder == '{CompanyName}':
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.bold = True
                else:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.bold = True

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        # Replace the placeholder with the value
                        original_text = cell.text
                        cell.text = original_text.replace(placeholder, str(value))

                        # Right-align the cells for InvoiceNo and InvoiceDate
                        if placeholder in ['{InvoiceNo}', '{InvoiceDate}']:
                            # Right alignment for specific placeholders
                            cell.paragraphs[0].alignment = 2  # 2 corresponds to right alignment

                        # Maintain original style without modification
                        for run in cell.paragraphs[0].runs:
                            if placeholder in run.text:
                                run.bold = True  # Bold the replaced text
                                run.italic = False  # Ensure it remains non-italic
                            else:
                                # Make sure non-placeholder text retains its original style
                                run.bold = False  # Set to false unless it's a placeholder

                        # If it's the Company Name, make it bold
                        if placeholder == '{CompanyName}':
                            for run in cell.paragraphs[0].runs:
                                if placeholder in run.text:
                                    run.bold = True  # Make Company Name bold

                        # For other placeholders
                        if placeholder == '{CompanyAddress}' or placeholder == '{ToWhom}':
                            for run in cell.paragraphs[0].runs:
                                if placeholder in run.text:
                                    run.bold = True



    # Assuming item_table is the second table in the document
    item_table = doc.tables[1]  # Adjust index if necessary

    # Clear existing item rows (keeping the header row)
    for row in item_table.rows[1:]:  # Start from the second row (index 1)
        for cell in row.cells:
            cell.text = ''  # Clear cell text

    # Insert item details into the item table starting from row 2
    for index, item in enumerate(items):
        if index + 1 < len(item_table.rows):  # Check if there's space for the item
            row_cells = item_table.rows[index + 1].cells  # Row 2 is index 1
        else:
            row_cells = item_table.add_row().cells  # Add new row if no space

        row_cells[0].text = str(item['description'])  # Description
        row_cells[1].text = str(item['quantity'])  # Quantity
        row_cells[2].text = f"{item['unit_price']:.2f}"  # Unit Price (MYR)
        row_cells[3].text = f"{item['total_price']:.2f}"  # Line Total (MYR)

    # Calculate total amount
    total_amount = sum(item['total_price'] for item in items)

    # Ensure there's a row for the total amount
    if len(item_table.rows) < 19:
        total_row = item_table.add_row().cells  # Create row 19 if not present
    else:
        total_row = item_table.rows[18].cells  # This is row 19 (0-indexed)

    # Set Total Amount label and value
    total_row[2].text = 'Total Amount'  # Label in column 3 (index 2)
    total_row[3].text = f'{total_amount:.2f}'  # Total Amount (MYR) in column 4 (index 3)

    # Make the Total Amount text bold
    for paragraph in total_row[3].paragraphs:
        for run in paragraph.runs:
            run.bold = True

    # Make the Total Amount label bold
    for paragraph in total_row[2].paragraphs:
        for run in paragraph.runs:
            run.bold = True


    # Save document with a unique invoice number
    filename = f'Invoice_{invoice_number}.docx'  # Use invoice_number in filename
    doc.save(filename)

    return filename


# Function to add a new company
def add_company():
    company_name = simpledialog.askstring("Input", "Enter Company Name:", parent=root)
    ssm_number = simpledialog.askstring("Input", "Enter SSM Number:", parent=root)
    address = simpledialog.askstring("Input", "Enter Address:", parent=root)
    phone = simpledialog.askstring("Input", "Enter Phone Number:", parent=root)
    
    if company_name and ssm_number:
        companies[company_name] = {
            'ssm_number': ssm_number,
            'address': address,
            'phone': phone
        }
        save_companies()  # Save companies to file
        company_list['values'] = list(companies.keys())
        messagebox.showinfo("Success", "Company added successfully!")
    else:
        messagebox.showerror("Error", "Please enter valid company details.")

# Function to delete selected item
def delete_item():
    selected_item = invoice_items.selection()
    if selected_item:
        invoice_items.delete(selected_item)
    else:
        messagebox.showwarning("Warning", "Please select an item to delete.")

def generate_invoice():
    company_name = company_list.get()
    if not company_name:
        messagebox.showerror("Error", "Please select a company.")
        return
    
    ssm_number = companies[company_name]['ssm_number']
    invoice_date = date_entry.get()

    # Get selected To Whom
    to_whom = to_whom_entry.get()
    if to_whom and to_whom not in to_whom_list:
        to_whom_list.append(to_whom)
        save_to_whom()  # Save updated To Whom list

    items = []
    for row in invoice_items.get_children():
        item = invoice_items.item(row)['values']
        if item:
            # Collecting details to create the items list
            quantity = int(item[0])  # First column is Quantity
            description = str(item[1])  # Ensure description is a string
            unit_price = float(item[2])  # Third column is Unit Price
            total_price = float(item[3])  # Fourth column is Line Total
            items.append({'description': description, 'quantity': quantity, 'unit_price': unit_price, 'total_price': total_price})
    
    if not items:
        messagebox.showerror("Error", "Please add at least one item.")
        return

    # Generate a simple invoice number
    existing_invoices = [int(f.split('_')[1].split('.')[0]) for f in os.listdir() if f.startswith('Invoice_') and f.endswith('.docx')]
    invoice_number = max(existing_invoices, default=0) + 1  # Increment the highest existing invoice number
    
    # Call create_invoice
    filename = create_invoice(company_name, ssm_number, invoice_number, invoice_date, items, to_whom)
    
    if messagebox.askyesno("Print Invoice", "Do you want to print the invoice?"):
        os.startfile(filename, "print")  # Will work on Windows

    messagebox.showinfo("Success", f"Invoice created successfully and saved as {filename}")

# Main GUI setup
root = tk.Tk()
root.title("E-Invoice Generator")

# Load existing companies and To Whom list
companies = load_companies()
to_whom_list = load_to_whom()

# Dropdown for company selection
tk.Label(root, text="Select Company:").grid(row=0, column=0)
company_list = ttk.Combobox(root)
company_list['values'] = list(companies.keys())
company_list.grid(row=0, column=1)

# Entry for "To Whom"
tk.Label(root, text="To Whom:").grid(row=1, column=0)
to_whom_entry = ttk.Combobox(root)
to_whom_entry['values'] = to_whom_list
to_whom_entry.grid(row=1, column=1)

# Button to add a new company
tk.Button(root, text="Add Company", command=add_company).grid(row=0, column=2)

# Entry for invoice date using DateEntry
tk.Label(root, text="Invoice Date:").grid(row=2, column=0)
date_entry = DateEntry(root, width=12, background='darkblue', foreground='white', borderwidth=2)
date_entry.grid(row=2, column=1)

# Treeview for invoice items
tk.Label(root, text="Invoice Items:").grid(row=3, column=0, columnspan=4)
invoice_items = ttk.Treeview(root, columns=("Quantity", "Description", "Unit Price", "Line Total"), show='headings')
invoice_items.heading("Quantity", text="Quantity")
invoice_items.heading("Description", text="Description")
invoice_items.heading("Unit Price", text="Unit Price (MYR)")
invoice_items.heading("Line Total", text="Line Total (MYR)")
invoice_items.grid(row=4, column=0, columnspan=4)

# Function to add item to invoice
def add_item():
    item_desc = simpledialog.askstring("Input", "Enter Item Description:", parent=root)
    item_quantity = simpledialog.askinteger("Input", "Enter Item Quantity:", parent=root)
    item_price = simpledialog.askfloat("Input", "Enter Item Unit Price (MYR):", parent=root)
    
    if item_desc and item_quantity is not None and item_price is not None:
        total_price = item_quantity * item_price  # Calculate total price for the item
        # Insert values as (Quantity, Description, Unit Price, Total Price)
        invoice_items.insert('', 'end', values=(item_quantity, item_desc, item_price, total_price))
    else:
        messagebox.showerror("Error", "Please enter valid item details.")

# Button to add item to the invoice
tk.Button(root, text="Add Item", command=add_item).grid(row=5, column=0)
# Button to delete item from the invoice
tk.Button(root, text="Delete Item", command=delete_item).grid(row=5, column=1)
# Button to generate invoice
tk.Button(root, text="Generate Invoice", command=generate_invoice).grid(row=6, column=0, columnspan=4)

# Run the application
root.mainloop()
